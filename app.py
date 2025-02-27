import os
import pandas as pd
import streamlit as st
from datetime import datetime
from docx import Document
from dotenv import load_dotenv
import requests
import re

# Load API key from .env file
load_dotenv()
API_KEY = os.getenv("NANONETS_API_KEY")
MODEL_ID = os.getenv("NANONETS_MODEL_ID")

# Function to convert PDF to JSON using NanoNets API
def convert_pdf_to_json(filepath):
    url = f'https://app.nanonets.com/api/v2/OCR/Model/{MODEL_ID}/LabelFile/'
    
    # Open the file and prepare the request data
    payload = {}

    files = [
        ('file', (os.path.basename(filepath), open(filepath, 'rb'), 'application/pdf'))
    ]

    # Prepare the headers with Authorization
    headers = {
        'accept': 'multipart/form-data'
        # 'Authorization': requests.auth.HTTPBasicAuth(API_KEY, '')
    }
        
    # Send POST request to the NanoNets API
    response = requests.request("POST", url, headers=headers, auth=requests.auth.HTTPBasicAuth(API_KEY, ''),data=payload, files=files)

    # Check for successful response and handle errors
    if response.status_code != 200:
        raise ValueError(f"API Error: {response.text}")

    # Parse and return JSON data
    print(response.text)
    return response.json()

# Function to extract tables from JSON and convert to CSV format
def convert_json_to_csv(json_data):
    tables = []
    
    # Extract tables from the JSON response
    for entry in json_data.get("result", []):
        for prediction in entry.get("prediction", []):
            if prediction.get("type") == "table":
                table_data = {}
                
                for cell in prediction.get("cells", []):
                    row, col = cell["row"], cell["col"]
                    text = cell["text"].strip(":").strip()
                    
                    if row not in table_data:
                        table_data[row] = {}
                    table_data[row][col] = text
                
                # Convert to DataFrame
                df = pd.DataFrame.from_dict(table_data, orient="index").sort_index()
                df = df.rename(columns=lambda x: f"Column_{x}")  # Standardize column names
                tables.append(df)

                
    # Merge all DataFrames in the list into one DataFrame
    if tables:
        combined_df = pd.concat(tables, ignore_index=True)
    else:
        raise ValueError("No tables found in the JSON response")
    
    return combined_df

# Function to process the file based on requirements
def process_file(df, filter_date):

    # Remove the first column
    df = df.iloc[:, 1:]

    # Remove rows with blank entries in the 'Column_6' column
    if 'Column_6' in df.columns:
        df = df[df['Column_6'].notna()]

    # Remove duplicate rows
    df = df.drop_duplicates()

    # Remove rows where the third column is blank
    if df.shape[1] > 2:  # Ensure the third column exists
        df = df[df.iloc[:, 2].notna()]

    # Promote the first row to header
    df.columns = df.iloc[0]
    df = df[1:]

    # Apply filters for the 'RESIGNATION' column
    resignation_col = next((col for col in df.columns if 'RESIGNATION' in col.upper()), None)

    if resignation_col:
        def convert_to_datetime(text):
            try:
                return datetime.strptime(text, "%d-%m-%Y")
            except ValueError:
                return None
            
        df["Extracted_Resignation_Date"] = df[resignation_col].apply(convert_to_datetime)
        
        filter_date = pd.to_datetime(filter_date)  # Ensure filter_date is datetime64[ns]
        df = df[(df["Extracted_Resignation_Date"].isna()) | (df["Extracted_Resignation_Date"] > filter_date)]
        
        df[resignation_col] = df["Extracted_Resignation_Date"].dt.strftime("%d-%m-%Y")
        df = df.drop(columns=["Extracted_Resignation_Date"])

    #Remove rows where the first column has 'COMPANIES'
    if df.shape[1] > 0:  # Ensure the first column exists
        df = df[~df.iloc[:, 0].astype(str).str.contains("COMPANIES", na=False, case=False)]

    return df

# Function to save the processed file as a Word document
def save_to_word(df, filepath):
    doc = Document()
    doc.add_heading("Processed Table", level=1)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for i, column in enumerate(df.columns):
        table.cell(0, i).text = column
    for row in df.itertuples(index=False):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value) if not pd.isna(value) else ""
    doc.save(filepath)

# Streamlit UI
st.title("SSM Table Extractor")

# Initialize session state
if "processed" not in st.session_state:
    st.session_state.processed = False
if "uploader_key" not in st.session_state:
    st.session_state.uploader_key = 0  # Initial key value for file uploader
if "selected_date" not in st.session_state:
    st.session_state.selected_date = datetime.now().date()

# User uploads file
uploaded_file = st.file_uploader("Upload a PDF file", type=["pdf"], key=f"uploader_{st.session_state.uploader_key}")

# User selects the year for filtering resignation dates
col1, col2 = st.columns([1, 3])
with col1:
    selected_date = st.date_input("Select a date:", value=st.session_state.selected_date, min_value=datetime(1900, 1, 1).date(), key=f"date_input_{st.session_state.uploader_key}")
actual_date = selected_date.replace(year=selected_date.year - 5)
st.write("Note: All entries with resignation date before this date will not be included.")

# Define temp file paths at the start
temp_pdf_path = "temp.pdf"
csv_path = "processed_data.csv"
word_path = "processed_data.docx"

if uploaded_file is not None:
    with open(temp_pdf_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    # Process the file when the user clicks the button
    if st.button("Process File"):
        try:
            json_data = convert_pdf_to_json(temp_pdf_path)
            df = convert_json_to_csv(json_data)
            df = process_file(df, actual_date)
            df.to_csv(csv_path, index=False)
            save_to_word(df, word_path)
            st.session_state.processed = True
        except Exception as e:
            st.error(f"Error: {e}")

# Display download buttons after processing
if st.session_state.processed:
    col1, col2 = st.columns(2)
    with col1:
        st.download_button("Download CSV", open(csv_path, "rb"), "processed_data.csv", "text/csv")
    with col2:
        st.download_button("Download Word Document", open(word_path, "rb"), "processed_data.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    
        
    # Cleanup: Remove temp files after the app runs
    if st.button("Process Another File"):
        try:
            # if os.path.exists(temp_pdf_path):
            #     os.remove(temp_pdf_path)
            # if os.path.exists(csv_path):
            #     os.remove(csv_path)
            # if os.path.exists(word_path):
            #     os.remove(word_path)
            st.session_state.processed = False
            st.session_state.uploader_key += 1  # Reset file uploader
            st.session_state.selected_date = datetime.now().date()  # Reset date input
            st.rerun()
        except Exception as e:
            st.warning(f"Could not delete temp files: {e}")